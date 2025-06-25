import os
import base64
import pandas as pd
from flask import jsonify, request, render_template, session
import requests
from werkzeug.utils import secure_filename
import json
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from PIL import Image
import io
import re  # 添加re模块导入
import PyPDF2  # 添加PDF处理库
import docx  # 添加Word文档处理库
import csv
import chardet  # 添加字符编码检测库
from datetime import datetime

# OpenRouter API配置
OPENROUTER_API_KEY = ""
OPENROUTER_API_URL = ""
OPENROUTER_MODEL = ""

# 创建一个带有重试机制的session
def create_retry_session():
    session = requests.Session()
    retry_strategy = Retry(
        total=3,  # 最多重试3次
        backoff_factor=1,  # 重试间隔时间
        status_forcelist=[429, 500, 502, 503, 504],  # 需要重试的HTTP状态码
    )
    adapter = HTTPAdapter(max_retries=retry_strategy)
    session.mount("http://", adapter)
    session.mount("https://", adapter)
    return session

def handle_ai_analysis(app):
    # 注册路由
    app.add_url_rule('/ai_analysis', 'ai_analysis_page', ai_analysis_page)
    app.add_url_rule('/api/ai_analysis', 'ai_analysis', ai_analysis, methods=['POST'])
    return app

def ai_analysis_page():
    return render_template('ai_analysis.html')

def process_image(file_path):
    """处理图片文件，转换为base64编码，并提取图片基本信息"""
    try:
        # 使用PIL打开图片并转换为RGB模式
        with Image.open(file_path) as img:
            # 获取图片基本信息
            img_info = {
                "format": img.format,
                "mode": img.mode,
                "width": img.width,
                "height": img.height,
            }
            
            # 转换为RGB模式（如果不是的话）
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # 调整图片大小以确保不超过API限制
            max_size = (1024, 1024)  # 增加最大尺寸以提高清晰度
            img.thumbnail(max_size, Image.Resampling.LANCZOS)
            
            # 将图片转换为JPEG格式的bytes
            buffer = io.BytesIO()
            img.save(buffer, format='JPEG', quality=90)  # 提高质量
            image_bytes = buffer.getvalue()
            
            # 转换为base64
            base64_image = base64.b64encode(image_bytes).decode('utf-8')
            
            return {
                "base64": base64_image,
                "info": img_info
            }
    except Exception as e:
        print(f"处理图片时出错: {str(e)}")
        return None

def extract_text_from_pdf(file_path):
    """从PDF文件中提取文本内容"""
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            # 提取文本（最多前10页，避免过大）
            max_pages = min(num_pages, 10)
            for page_num in range(max_pages):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n\n"
            
            # 如果有更多页，添加提示
            if num_pages > max_pages:
                text += f"\n[注意：此PDF共有{num_pages}页，仅显示前{max_pages}页内容]\n"
                
        return {
            "text": text,
            "pages": num_pages
        }
    except Exception as e:
        print(f"提取PDF文本时出错: {str(e)}")
        return {"text": f"无法提取PDF内容: {str(e)}", "pages": 0}

def extract_text_from_docx(file_path):
    """从Word文档中提取文本内容"""
    try:
        doc = docx.Document(file_path)
        text = ""
        
        # 提取段落文本
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                text += " | ".join(row_text) + "\n"
            text += "\n"
            
        return text
    except Exception as e:
        print(f"提取Word文档文本时出错: {str(e)}")
        return f"无法提取Word文档内容: {str(e)}"

def extract_text_from_txt(file_path):
    """从文本文件中提取文本内容"""
    try:
        # 尝试检测文件编码
        try:
            import chardet
            with open(file_path, 'rb') as f:
                raw_data = f.read(4096)  # 读取前4KB来检测编码
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'utf-8'
                confidence = result.get('confidence', 0)
                print(f"检测到文件编码: {encoding}，置信度: {confidence}")
        except Exception as e:
            print(f"检测文件编码失败: {str(e)}，使用默认UTF-8")
            encoding = 'utf-8'  # 如果检测失败，默认使用UTF-8
        
        # 读取文件内容，使用'replace'错误处理模式以避免因编码问题导致的失败
        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
            content = f.read(100000)  # 限制读取的内容大小，避免过大
            
            # 如果内容太长，截断并添加说明
            if len(content) > 100000:
                content = content[:100000] + "\n\n[... 内容过长，已截断 ...]"
                
            # 处理不规则的换行和空白
            content = content.replace('\r\n', '\n').replace('\r', '\n')
            
            return content
    except Exception as e:
        print(f"读取文本文件 {file_path} 失败: {str(e)}")
        return f"[无法读取文件内容: {str(e)}]"

def extract_text_from_csv(file_path):
    """从CSV文件中提取内容，自动检测编码"""
    try:
        # 检测文件编码
        with open(file_path, 'rb') as f:
            raw_data = f.read(4096)  # 读取前4KB来检测编码
            result = chardet.detect(raw_data)
            encoding = result['encoding'] or 'utf-8'
        
        # 使用检测到的编码读取CSV
        rows = []
        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
            csv_reader = csv.reader(f)
            for i, row in enumerate(csv_reader):
                rows.append(row)
                if i >= 100:  # 最多读取100行
                    break
        
        # 格式化为表格形式的文本
        if not rows:
            return "CSV文件为空"
        
        text = ""
        for row in rows:
            text += " | ".join(row) + "\n"
        
        return text
    except Exception as e:
        print(f"提取CSV文件内容时出错: {str(e)}")
        return f"无法提取CSV文件内容: {str(e)}"

def record_ai_analysis(username, file_count, response_time=0):
    """记录AI分析操作"""
    try:
        from app import get_db, logger
        with get_db() as conn:
            cursor = conn.cursor()
            # 检查表是否存在
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS ai_analysis_records (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    username TEXT NOT NULL,
                    file_count INTEGER NOT NULL,
                    response_time REAL NOT NULL,
                    status TEXT NOT NULL,
                    created_at DATETIME NOT NULL
                )
            ''')
            
            # 插入记录
            cursor.execute('''
                INSERT INTO ai_analysis_records (
                    username, file_count, response_time, status, created_at
                ) VALUES (?, ?, ?, ?, ?)
            ''', (
                username,
                file_count,
                response_time,
                '完成',
                datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            ))
            conn.commit()
            
            logger.info(f'记录AI分析操作 - 用户: {username}, 文件数: {file_count}, 响应时间: {response_time}秒')
            
    except Exception as e:
        print(f"记录AI分析操作时出错: {str(e)}")  # 调试信息

def admin_ai_stats():
    """获取AI分析统计数据"""
    try:
        from app import get_db
        with get_db() as conn:
            cursor = conn.cursor()
            
            # 获取今日AI分析次数
            today = datetime.now().strftime('%Y-%m-%d')
            cursor.execute('''
                SELECT COUNT(*) as count 
                FROM ai_analysis_records 
                WHERE date(created_at) = ?
            ''', (today,))
            today_count = cursor.fetchone()['count']
            
            # 获取总AI分析次数
            cursor.execute('SELECT COUNT(*) as count FROM ai_analysis_records')
            total_count = cursor.fetchone()['count']
            
            # 获取使用过AI分析的用户数
            cursor.execute('SELECT COUNT(DISTINCT username) as count FROM ai_analysis_records')
            ai_users = cursor.fetchone()['count']
            
            # 获取平均响应时间
            cursor.execute('SELECT AVG(response_time) as avg_time FROM ai_analysis_records')
            result = cursor.fetchone()
            avg_response_time = round(result['avg_time'], 2) if result['avg_time'] is not None else 0
            
            # 获取最近10条AI分析记录
            cursor.execute('''
                SELECT username, file_count, response_time, status, created_at
                FROM ai_analysis_records 
                ORDER BY created_at DESC 
                LIMIT 10
            ''')
            recent_ai_analysis = []
            for record in cursor.fetchall():
                recent_ai_analysis.append({
                    'username': record['username'],
                    'file_count': record['file_count'],
                    'response_time': record['response_time'],
                    'status': record['status'],
                    'created_at': record['created_at']
                })
            
            stats = {
                'today_ai_analysis': today_count,
                'total_ai_analysis': total_count,
                'ai_users': ai_users,
                'avg_response_time': avg_response_time,
                'recent_ai_analysis': recent_ai_analysis
            }
            
            return jsonify(stats)
            
    except Exception as e:
        print(f"获取AI分析统计数据时出错: {str(e)}")  # 调试信息
        return jsonify({'error': '获取AI统计数据失败'}), 500

def ai_analysis():
    try:
        # 记录开始时间
        start_time = datetime.now()
        
        # 获取用户消息
        message = request.form.get('message', '')
        
        # 检查是否是编辑的消息
        is_edit_message = request.form.get('is_edit_message') == 'true'
        original_filenames = request.form.get('original_filenames', '')
        
        # 详细记录请求信息
        #print("=" * 50)
        #print("收到AI分析请求:")
        #print(f"消息内容: {message[:100]}..." if len(message) > 100 else f"消息内容: {message}")
        #print(f"是否是编辑消息: {is_edit_message}")
        #print(f"原始文件名: {original_filenames}")
        #print("表单数据:", list(request.form.keys()))
        #print("=" * 50)
        
        # 获取上传的文件
        files = request.files.getlist('files[]')
        data_description = ""
        image_data = []
        file_content_texts = []  # 统一使用file_content_texts
        
        # 处理上传的文件
        if files:
            print(f"收到 {len(files)} 个文件")  # 调试日志
            print("文件名列表:", [f.filename for f in files])
            dataframes = []
            file_descriptions = []
            
            # 确保uploads目录存在
            uploads_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
            os.makedirs(uploads_dir, exist_ok=True)
            
            for file in files:
                if not file.filename:
                    continue
                    
                print(f"处理文件: {file.filename}")  # 调试日志
                # 保留原始文件名，但确保它是安全的
                filename = secure_filename(file.filename)
                # 将文件保存到uploads目录而不是temp目录
                file_path = os.path.join(uploads_dir, filename)
                
                try:
                    # 保存文件
                    file.save(file_path)
                    print(f"文件已保存到: {file_path}")  # 调试日志
                    
                    # 获取文件扩展名
                    ext = os.path.splitext(filename)[1].lower() if '.' in filename else ''
                    
                    # 根据文件类型处理
                    if ext in ['.xlsx', '.xls']:
                        try:
                            df = pd.read_excel(file_path)
                            dataframes.append({
                                'name': filename,
                                'data': df
                            })
                            file_descriptions.append(f"Excel文件 '{filename}' 已上传，包含 {len(df)} 行数据。")
                        except Exception as e:
                            file_descriptions.append(f"Excel文件 '{filename}' 读取失败：{str(e)}")
                    
                    elif ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
                        # 处理图片文件
                        img_data = process_image(file_path)
                        if img_data:
                            image_data.append({
                                'name': filename,
                                'data': img_data['base64'],
                                'path': file_path  # 添加文件路径以便后续引用
                            })
                            # 添加图片信息到描述
                            info = img_data['info']
                            file_descriptions.append(f"图片文件 '{filename}' 已上传。尺寸: {info['width']}x{info['height']}，格式: {info['format']}。")
                        else:
                            file_descriptions.append(f"图片文件 '{filename}' 处理失败。")
                    
                    elif ext == '.pdf':
                        # 处理PDF文件
                        pdf_data = extract_text_from_pdf(file_path)
                        if pdf_data:
                            file_descriptions.append(f"PDF文件 '{filename}' 已上传，共 {pdf_data['pages']} 页。")
                            file_content_texts.append(f"文件名: {filename}\n类型: pdf\n内容:\n{pdf_data['text']}")
                            print(f"成功处理PDF文件: {filename}")
                        else:
                            file_descriptions.append(f"PDF文件 '{filename}' 处理失败。")
                    
                    elif ext in ['.doc', '.docx']:
                        # 处理Word文档
                        doc_text = extract_text_from_docx(file_path)
                        file_descriptions.append(f"Word文档 '{filename}' 已上传。")
                        file_content_texts.append(f"文件名: {filename}\n类型: doc\n内容:\n{doc_text}")
                        print(f"成功处理Word文档: {filename}")
                    
                    elif ext == '.txt' or filename == 'requirements.txt' or filename == 'README.md':
                        # 处理文本文件 (保留对特殊文件的支持)
                        txt_content = extract_text_from_txt(file_path)
                        file_descriptions.append(f"文本文件 '{filename}' 已上传。")
                        file_content_texts.append(f"文件名: {filename}\n类型: txt\n内容:\n{txt_content}")
                        print(f"成功处理文本文件: {filename}")
                    
                    elif ext == '.csv':
                        # 处理CSV文件
                        csv_content = extract_text_from_csv(file_path)
                        file_descriptions.append(f"CSV文件 '{filename}' 已上传。")
                        file_content_texts.append(f"文件名: {filename}\n类型: csv\n内容:\n{csv_content}")
                        print(f"成功处理CSV文件: {filename}")
                    
                    else:
                        # 对于其他类型的文件，尝试读取内容（如果是文本文件）
                        try:
                            # 检测文件是否为文本
                            is_text = True
                            with open(file_path, 'rb') as f:
                                chunk = f.read(1024)
                                if b'\x00' in chunk:  # 包含空字节表示是二进制文件
                                    is_text = False
                            
                            if is_text:
                                # 尝试检测编码
                                try:
                                    import chardet
                                    with open(file_path, 'rb') as f:
                                        raw_data = f.read(4096)  # 读取前4KB来检测编码
                                        result = chardet.detect(raw_data)
                                        encoding = result['encoding'] or 'utf-8'
                                except:
                                    encoding = 'utf-8'  # 如果检测失败，默认使用UTF-8
                                
                                # 读取文件内容
                                with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                                    file_content = f.read(50000)  # 限制读取的内容大小，避免过大
                                    # 如果内容太长，截断并添加说明
                                    if len(file_content) > 50000:
                                        file_content = file_content[:50000] + "\n\n[... 内容过长，已截断 ...]"
                                    file_descriptions.append(f"文件 '{filename}' 已上传并成功读取内容。")
                                    file_content_texts.append(f"文件名: {filename}\n类型: 文本\n内容:\n{file_content}")
                                    print(f"成功处理通用文本文件: {filename}")
                            else:
                                # 二进制文件，只添加基本信息
                                file_size = os.path.getsize(file_path)
                                size_str = f"{file_size/1024/1024:.1f}MB" if file_size > 1024*1024 else f"{file_size/1024:.1f}KB"
                                file_descriptions.append(f"文件 '{filename}' ({size_str}) 已上传。")
                        except Exception as e:
                            # 如果读取失败，只添加基本信息
                            file_size = os.path.getsize(file_path)
                            size_str = f"{file_size/1024/1024:.1f}MB" if file_size > 1024*1024 else f"{file_size/1024:.1f}KB"
                            file_descriptions.append(f"文件 '{filename}' ({size_str}) 已上传。")
                            print(f"尝试读取文件 {filename} 内容时出错: {str(e)}")  # 调试日志
                
                except Exception as e:
                    print(f"处理文件 {filename} 时出错: {str(e)}")  # 调试日志
                    file_descriptions.append(f"处理文件 '{filename}' 时出错: {str(e)}")
                
                # 移除finally块，不再删除上传的文件，使其可用于预览
            
            # 准备文件数据描述
            if dataframes:
                data_description = prepare_data_description(dataframes)
            
            # 添加文件处理结果到消息中
            if file_descriptions:
                message = f"{message}\n\n文件信息：\n" + "\n".join(file_descriptions)
                print(f"添加了 {len(file_descriptions)} 条文件描述")
            
            # 添加文件内容到消息中
            if file_content_texts:
                message += "\n\n文件内容：\n" + "\n\n---\n\n".join(file_content_texts)
                print(f"添加了 {len(file_content_texts)} 个文件的内容")
                for i, text in enumerate(file_content_texts):
                    print(f"文件内容 {i+1} 的前100个字符: {text[:100]}...")
        
        # 处理编辑消息但没有收到文件的情况
        elif is_edit_message and original_filenames:
            print("这是编辑的消息，尝试查找已上传的文件")
            print(f"需要查找的文件名: {original_filenames}")
            file_names = original_filenames.split(',')
            dataframes = []
            file_descriptions = []
            file_content_texts = []
            
            # 确保uploads目录存在
            uploads_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
            
            if os.path.exists(uploads_dir):
                print(f"查找文件夹中的文件: {uploads_dir}")
                # 列出uploads目录中的所有文件
                all_files = os.listdir(uploads_dir)
                print(f"目录中找到 {len(all_files)} 个文件:")
                for idx, f in enumerate(all_files):
                    print(f"  {idx+1}. {f}")
                
                # 尝试找到匹配的文件
                for file_name in file_names:
                    file_name = file_name.strip()
                    found = False
                    matched_file = None
                    
                    # 首先尝试完全匹配
                    if file_name in all_files:
                        matched_file = file_name
                        found = True
                        print(f"找到完全匹配: {file_name}")
                    else:
                        # 尝试部分匹配
                        for existing_file in all_files:
                            if file_name in existing_file or existing_file in file_name:
                                matched_file = existing_file
                                found = True
                                print(f"找到部分匹配: {existing_file} (原名: {file_name})")
                                break
                    
                    if found and matched_file:
                        file_path = os.path.join(uploads_dir, matched_file)
                        print(f"处理找到的文件: {file_path}")
                        
                        try:
                            # 获取文件扩展名
                            ext = os.path.splitext(file_path)[1].lower()
                            print(f"文件扩展名: {ext}")
                            
                            # 根据文件类型处理
                            if ext in ['.xlsx', '.xls']:
                                try:
                                    df = pd.read_excel(file_path)
                                    dataframes.append({
                                        'name': os.path.basename(file_path),
                                        'data': df
                                    })
                                    file_descriptions.append(f"Excel文件 '{os.path.basename(file_path)}' 已找到，包含 {len(df)} 行数据。")
                                    print(f"成功处理Excel文件: {file_path}")
                                except Exception as e:
                                    file_descriptions.append(f"Excel文件 '{os.path.basename(file_path)}' 读取失败：{str(e)}")
                                    print(f"处理Excel文件失败: {str(e)}")
                            
                            elif ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
                                # 处理图片文件
                                img_data = process_image(file_path)
                                if img_data:
                                    image_data.append({
                                        'name': os.path.basename(file_path),
                                        'data': img_data['base64'],
                                        'path': file_path  # 添加文件路径以便后续引用
                                    })
                                    # 添加图片信息到描述
                                    info = img_data['info']
                                    file_descriptions.append(f"图片文件 '{os.path.basename(file_path)}' 已找到。尺寸: {info['width']}x{info['height']}，格式: {info['format']}。")
                                    print(f"成功处理图片文件: {file_path}")
                                else:
                                    file_descriptions.append(f"图片文件 '{os.path.basename(file_path)}' 处理失败。")
                                    print(f"处理图片文件失败")
                            
                            elif ext == '.pdf':
                                # 处理PDF文件
                                pdf_data = extract_text_from_pdf(file_path)
                                if pdf_data:
                                    file_content_texts.append(f"文件名: {os.path.basename(file_path)}\n类型: pdf\n内容:\n{pdf_data['text']}")
                                    file_descriptions.append(f"PDF文件 '{os.path.basename(file_path)}' 已找到，共 {pdf_data['pages']} 页。")
                                    print(f"成功处理PDF文件: {file_path}")
                                else:
                                    file_descriptions.append(f"PDF文件 '{os.path.basename(file_path)}' 处理失败。")
                                    print(f"处理PDF文件失败")
                            
                            elif ext == '.docx' or ext == '.doc':
                                # 处理Word文档
                                text = extract_text_from_docx(file_path)
                                file_content_texts.append(f"文件名: {os.path.basename(file_path)}\n类型: doc\n内容:\n{text}")
                                file_descriptions.append(f"Word文档 '{os.path.basename(file_path)}' 已找到，提取了 {len(text.split())} 个单词。")
                                print(f"成功处理Word文件: {file_path}")
                            
                            elif ext == '.txt':
                                # 处理文本文件
                                text = extract_text_from_txt(file_path)
                                file_content_texts.append(f"文件名: {os.path.basename(file_path)}\n类型: txt\n内容:\n{text}")
                                file_descriptions.append(f"文本文件 '{os.path.basename(file_path)}' 已找到，提取了 {len(text.split())} 个单词。")
                                print(f"成功处理文本文件: {file_path}")
                            
                            elif ext == '.csv':
                                # 处理CSV文件
                                try:
                                    df, text = extract_text_from_csv(file_path)
                                    dataframes.append({
                                        'name': os.path.basename(file_path),
                                        'data': df
                                    })
                                    file_content_texts.append(f"文件名: {os.path.basename(file_path)}\n类型: csv\n内容:\n{text}")
                                    file_descriptions.append(f"CSV文件 '{os.path.basename(file_path)}' 已找到，包含 {len(df)} 行数据。")
                                    print(f"成功处理CSV文件: {file_path}")
                                except Exception as e:
                                    file_descriptions.append(f"CSV文件 '{os.path.basename(file_path)}' 读取失败：{str(e)}")
                                    print(f"处理CSV文件失败: {str(e)}")
                            
                            else:
                                # 尝试作为通用文件处理
                                try:
                                    with open(file_path, 'rb') as f:
                                        is_binary = False
                                        try:
                                            for block in f:
                                                if b'\0' in block:
                                                    is_binary = True
                                                    break
                                        except:
                                            is_binary = True
                                    
                                    if not is_binary:
                                        # 文本文件，尝试读取内容
                                        try:
                                            import chardet
                                            with open(file_path, 'rb') as f:
                                                raw_data = f.read(4096)  # 读取前4KB来检测编码
                                                result = chardet.detect(raw_data)
                                                encoding = result['encoding'] or 'utf-8'
                                        except:
                                            encoding = 'utf-8'  # 如果检测失败，默认使用UTF-8
                                        
                                        # 读取文件内容
                                        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                                            file_content = f.read(50000)  # 限制读取的内容大小，避免过大
                                            # 如果内容太长，截断并添加说明
                                            if len(file_content) > 50000:
                                                file_content = file_content[:50000] + "\n\n[... 内容过长，已截断 ...]"
                                            file_content_texts.append(f"文件名: {os.path.basename(file_path)}\n类型: 文本\n内容:\n{file_content}")
                                            file_descriptions.append(f"文件 '{os.path.basename(file_path)}' 已找到并成功读取内容。")
                                            print(f"成功作为文本文件处理: {file_path}")
                                    else:
                                        # 二进制文件，只添加基本信息
                                        file_size = os.path.getsize(file_path)
                                        size_str = f"{file_size/1024/1024:.1f}MB" if file_size > 1024*1024 else f"{file_size/1024:.1f}KB"
                                        file_descriptions.append(f"文件 '{os.path.basename(file_path)}' ({size_str}) 已找到，但无法读取二进制内容。")
                                        print(f"文件是二进制格式: {file_path}")
                                except Exception as e:
                                    file_descriptions.append(f"不支持的文件类型: {os.path.basename(file_path)}")
                                    print(f"不支持的文件类型: {ext}, 错误: {str(e)}")
                        except Exception as e:
                            print(f"处理文件时出错: {str(e)}")
                            file_descriptions.append(f"处理文件 '{os.path.basename(file_path)}' 时出错: {str(e)}")
                    else:
                        print(f"未找到匹配的文件: {file_name}")
                        file_descriptions.append(f"未找到文件 '{file_name}'")
                
                # 处理找到的DataFrame数据
                if dataframes:
                    data_description = prepare_data_description(dataframes)
                
                # 添加文件处理结果到消息中
                if file_descriptions:
                    message = f"{message}\n\n文件信息：\n" + "\n".join(file_descriptions)
                    print(f"添加了 {len(file_descriptions)} 条文件描述")
                
                # 添加文件内容到消息中
                if file_content_texts:
                    message += "\n\n文件内容：\n" + "\n\n---\n\n".join(file_content_texts)
                    print(f"添加了 {len(file_content_texts)} 个文件的内容")
                    for i, text in enumerate(file_content_texts):
                        print(f"文件内容 {i+1} 的前100个字符: {text[:100]}...")
            else:
                print(f"上传目录不存在: {uploads_dir}")
                file_descriptions.append("上传目录不存在，无法找到之前上传的文件")
                
            # 处理找到的DataFrame数据
            if dataframes:
                data_description = prepare_data_description(dataframes)
                # 将数据文件描述添加到请求中
                message += "\n\n数据文件：\n" + "\n".join(file_descriptions)
        
        # 如果没有消息也没有文件，返回错误
        if not message and not data_description and not image_data:
            return jsonify({'success': False, 'message': '请输入问题或上传文件'})

        # 调用OpenRouter API进行分析
        response = get_deepseek_analysis(message, data_description, image_data)
        
        # 计算响应时间（秒）
        response_time = (datetime.now() - start_time).total_seconds()
        
        # 记录分析操作
        username = session.get('username', '匿名用户')
        file_count = len(files) if files else 0
        record_ai_analysis(username, file_count, response_time)

        return jsonify({
            'success': True,
            'response': response
        })

    except Exception as e:
        print(f"处理请求时出错: {str(e)}")  # 调试日志
        return jsonify({
            'success': False,
            'message': f'处理请求时出现错误：{str(e)}'
        })

def allowed_file(filename):
    """检查文件类型是否允许"""
    ALLOWED_EXTENSIONS = {
        # 文档类型
        'xlsx', 'xls', 'csv', 'txt', 'pdf', 'doc', 'docx', 'ppt', 'pptx', 'md', 'rtf',
        # 图像类型
        'jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp', 'svg', 'tiff', 'tif',
        # 开发文件
        'py', 'js', 'html', 'css', 'java', 'c', 'cpp', 'h', 'php', 'sql', 'json', 'xml', 'yaml', 'yml',
        # 配置文件
        'ini', 'cfg', 'conf', 'log', 'sh', 'bat', 'ps1',
        # 其他编程语言
        'rb', 'pl', 'go', 'ts', 'jsx', 'tsx', 'vue', 'dart', 'swift', 'kt', 'rs', 'scala', 'lua', 'r'
    }
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def prepare_data_description(dataframes):
    """准备数据描述信息"""
    description = []
    
    for df_info in dataframes:
        df = df_info['data']
        desc = f"文件 '{df_info['name']}' 包含以下信息：\n"
        desc += f"- 总行数：{len(df)}\n"
        desc += f"- 列名：{', '.join(df.columns.tolist())}\n"
        
        # 添加每列的基本统计信息
        for column in df.columns:
            if pd.api.types.is_numeric_dtype(df[column]):
                stats = df[column].describe()
                desc += f"\n{column} 列的统计信息：\n"
                desc += f"- 平均值：{stats['mean']:.2f}\n"
                desc += f"- 最小值：{stats['min']:.2f}\n"
                desc += f"- 最大值：{stats['max']:.2f}\n"
            else:
                unique_values = df[column].nunique()
                desc += f"\n{column} 列包含 {unique_values} 个不同的值\n"
        
        description.append(desc)
    
    return "\n\n".join(description)

def get_deepseek_analysis(user_message, data_description="", image_data=None):
    """调用OpenRouter API进行分析"""
    try:
        if not OPENROUTER_API_KEY:
            raise Exception("未设置OpenRouter API密钥")

        # 设置最大上下文长度
        MAX_CONTEXT_LENGTH = 2000  # 根据模型token限制设置

        # 构建系统提示信息
        system_prompt = """你是一个专业的AI助手,可以:
1. 回答用户的一般性问题
2. 分析Excel数据并提供见解
3. 分析和描述图片内容
4. 处理和分析上传的文件

当分析数据、图片或文件时:
1. 保持专业性和准确性
2. 使用清晰的语言解释
3. 详细描述图片内容
4. 如果可能,提供具体的数字支持
5. 如果数据不足以回答问题,请明确指出
6. 如果需要更多信息,请具体说明需要什么数据"""

        # 准备用户消息
        messages = [
            {"role": "system", "content": system_prompt}
        ]

        # 添加历史消息到消息列表中
        try:
            history = request.form.get('history', '[]')
            if history:
                history_messages = json.loads(history)
                # 确保按时间顺序添加所有历史消息
                for msg in history_messages:
                    # 确保消息类型正确
                    if msg.get('type') not in ['user', 'ai']:
                        continue
                    role = "assistant" if msg['type'] == 'ai' else "user"
                    content = msg.get('content', '')
                    # 清理HTML标签
                    content = re.sub(r'<[^>]+>', '', content)
                    messages.append({"role": role, "content": content.strip()})
        except Exception as e:
            print(f"处理历史消息时出错: {str(e)}")

        # 如果有图片数据，添加到消息中
        if image_data:
            for img in image_data:
                messages.append({
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{img['data']}"
                            }
                        }
                    ]
                })

        # 添加用户消息和数据描述
        user_content = user_message
        if data_description:
            user_content = f"{user_message}\n\n{data_description}"
        
        # 打印详细的消息内容日志
        #print("=" * 50)
        #print("发送给AI的最终消息内容:")
        #print("-" * 30)
        #print(user_content[:1000] + "..." if len(user_content) > 1000 else user_content)
        #print("=" * 50)
        
        messages.append({"role": "user", "content": user_content})

        # 控制上下文长度
        current_length = sum(len(str(m.get('content', ''))) for m in messages)
        while current_length > MAX_CONTEXT_LENGTH and len(messages) > 2:  # 保留system和最新的用户消息
            removed = messages.pop(1)  # 移除最早的历史消息
            current_length = sum(len(str(m.get('content', ''))) for m in messages)
            print(f"移除历史消息以控制上下文长度，当前长度：{current_length}")

        # 准备请求数据
        headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json",
            "HTTP-Referer": "http://localhost:5001",
            "X-Title": "Excel Data Analyzer"
        }

        data = {
            "model": OPENROUTER_MODEL,
            "messages": messages,
            "temperature": 0.7,
            "max_tokens": 2000
        }

        #print(f"发送请求到OpenRouter API...")  # 调试日志

        # 使用带有重试机制的session发送请求
        session = create_retry_session()
        try:
            response = session.post(
                OPENROUTER_API_URL, 
                headers=headers, 
                json=data,
                timeout=(10, 90)
            )
            
            #print(f"API响应状态码: {response.status_code}")  # 调试日志
            
            # 检查响应状态码
            if response.status_code != 200:
                print(f"API错误响应: {response.text}")  # 调试日志
                return f"API请求失败，状态码：{response.status_code}，错误信息：{response.text}"
            
            # 尝试解析JSON响应
            try:
                result = response.json()
            except json.JSONDecodeError as e:
                print(f"JSON解析错误: {str(e)}")  # 调试日志
                print(f"原始响应内容: {response.text}")  # 调试日志
                return "API返回了无效的JSON响应"

            # 检查响应格式
            if "choices" not in result or not result["choices"]:
                print(f"无效的API响应格式: {result}")  # 调试日志
                return "API返回了无效的响应格式"

            # 获取响应内容
            content = result["choices"][0].get("message", {}).get("content")
            if not content:
                print(f"响应中没有找到内容: {result}")  # 调试日志
                return "API响应中没有找到有效的内容"

            return content

        except requests.exceptions.Timeout:
            print("API请求超时")  # 调试日志
            return "API请求超时，请稍后重试"
        except requests.exceptions.RequestException as e:
            print(f"API请求异常: {str(e)}")  # 调试日志
            return f"API请求失败: {str(e)}"

    except Exception as e:
        print(f"处理错误详情: {str(e)}")  # 调试日志
        return f"处理请求时出现错误：{str(e)}" 
